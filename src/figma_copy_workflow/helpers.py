"""Helper functions for Figma Copy Workflow."""

import csv
import re
from collections import defaultdict
from typing import Dict, List

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn


def read_csv_data(csv_file_path: str) -> List[Dict[str, str]]:
    """Read CSV data and return as list of dictionaries."""
    data = []
    with open(csv_file_path, 'r', encoding='utf-8') as file:
        # Handle potential BOM and read content
        content = file.read()
        if content.startswith('\ufeff'):
            content = content[1:]
        
        # Split into lines and process
        lines = content.strip().split('\n')
        reader = csv.DictReader(lines)
        
        for row in reader:
            # Clean up the data - remove leading/trailing whitespace and tabs
            cleaned_row = {}
            for key, value in row.items():
                cleaned_key = key.strip(' \t"')
                cleaned_value = value.strip(' \t"') if value else ''
                cleaned_row[cleaned_key] = cleaned_value
            data.append(cleaned_row)
    
    return data


def group_data_by_section(data: List[Dict[str, str]]) -> Dict[str, List[Dict[str, str]]]:
    """Group data by the 'group' column."""
    grouped_data = defaultdict(list)
    
    for row in data:
        group = row.get('group', 'Unknown Group')
        if group and group.strip():  # Only add rows with non-empty groups
            grouped_data[group].append(row)
    
    return dict(grouped_data)


def set_cell_background_color(cell, color_rgb):
    """Set background color for a table cell."""
    cell_xml_element = cell._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), color_rgb)
    table_cell_properties.append(shade_obj)


def create_word_document(grouped_data: Dict[str, List[Dict[str, str]]], output_path: str) -> None:
    """Create a Word document with headers by section and tables for each group."""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Figma Copy Export', 0)
    # Make title gray
    for run in title.runs:
        run.font.color.rgb = RGBColor(105, 105, 105)  # Dark gray
    # Reduce spacing after title
    title.paragraph_format.space_after = Inches(0.1)
    
    for group_name, rows in grouped_data.items():
        if not rows:  # Skip empty groups
            continue
            
        # Add section header
        heading = doc.add_heading(group_name, level=1)
        # Make heading gray
        for run in heading.runs:
            run.font.color.rgb = RGBColor(105, 105, 105)  # Dark gray
        
        # Reduce spacing before and after heading
        heading.paragraph_format.space_before = Inches(0.1)
        heading.paragraph_format.space_after = Inches(0.05)
        
        # Create table with headers: Label, Text, ID
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Adjust column widths to fit within page margins (total ~6.5" usable width)
        table.columns[0].width = Inches(1.5)    # Label - compact
        table.columns[1].width = Inches(4.0)    # Text - most space for content
        table.columns[2].width = Inches(1.0)    # ID - minimal space
        
        # Add table headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Label'
        header_cells[1].text = 'Text'
        header_cells[2].text = 'ID'
        
        # Format headers: bold and light gray background for Label and ID columns
        for i, cell in enumerate(header_cells):
            # Make all headers bold and gray text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(105, 105, 105)  # Dark gray text
                # Reduce font size slightly for better fit
                paragraph.runs[0].font.size = Inches(0.12)  # ~8.5pt
            
            # Set light gray background for Label (0) and ID (2) columns
            if i != 1:  # Not the Text column
                set_cell_background_color(cell, 'D3D3D3')  # Light gray
            
            # Set vertical alignment
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row_data.get('layer_name', '')
            row_cells[1].text = row_data.get('figma_text', '')
            row_cells[2].text = row_data.get('id', '')
            
            # Set light gray background for Label (0) and ID (2) columns
            for i, cell in enumerate(row_cells):
                if i != 1:  # Not the Text column
                    set_cell_background_color(cell, 'D3D3D3')  # Light gray
                
                # Set vertical alignment and allow text wrapping
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                # Enable text wrapping for better content display
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 0  # Left align
                    # Reduce font size slightly for better fit
                    for run in paragraph.runs:
                        run.font.size = Inches(0.11)  # ~8pt
        
        # Add minimal space after each table instead of a full paragraph
        space_para = doc.add_paragraph()
        space_para.paragraph_format.space_before = Inches(0)
        space_para.paragraph_format.space_after = Inches(0.05)
    
    # Save the document
    doc.save(output_path)


def read_word_document_data(word_file_path: str, preserve_formatting: bool = True) -> Dict[str, str]:
    """Read Word document and extract updated text content mapped by ID.
    
    Args:
        word_file_path: Path to the input Word document
        preserve_formatting: Whether to preserve formatting as Markdown (True) or extract plain text (False)
        
    Returns:
        Dictionary mapping ID to updated text content
    """
    doc = Document(word_file_path)
    id_to_text = {}
    
    def _extract_formatted_text_from_cell(cell) -> str:
        """
        Extract text from a cell while preserving formatting.
        
        Args:
            cell: Cell object from python-docx
            
        Returns:
            Text with Markdown formatting preserved or plain text based on preserve_formatting setting
        """
        # If formatting preservation is disabled, return plain text
        if not preserve_formatting:
            return cell.text.strip()
        
        formatted_text = ""
        
        for para in cell.paragraphs:
            para_text = ""
            
            # Check if this paragraph is a list item
            is_list_item = False
            is_numbered_list = False
            list_marker = ""
            if para.text.strip():
                text = para.text.strip()
                # Check for unordered list markers
                if text.startswith(('•', '-', '*')):
                    is_list_item = True
                    list_marker = text[0]
                # Check for numbered lists (more comprehensive)
                elif len(text) > 2:
                    # Check for patterns like "1.", "2)", "a.", "A)", "i.", "IV.", etc.
                    numbered_pattern = r'^(\d+[.)]|[a-zA-Z][.)]|[ivxlcdm]+[.)]|[IVXLCDM]+[.)])\s'
                    match = re.match(numbered_pattern, text, re.IGNORECASE)
                    if match:
                        is_list_item = True
                        is_numbered_list = True
                        list_marker = match.group(1)
            
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                
                # Check if this run contains a hyperlink
                hyperlink_url = None
                
                # Method 1: Check if run's parent is a hyperlink element
                parent = run.element.getparent()
                while parent is not None:
                    if parent.tag.endswith('hyperlink'):
                        # Found hyperlink element, extract the relationship ID
                        rel_id = parent.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rel_id:
                            try:
                                hyperlink_url = doc.part.rels[rel_id].target_ref
                                break
                            except (KeyError, AttributeError):
                                pass  # Invalid relationship, ignore hyperlink
                    parent = parent.getparent()
                
                # Method 2: Check for hyperlinks in the paragraph's XML structure
                if not hyperlink_url:
                    for hyperlink in para.element.iter():
                        if hyperlink.tag.endswith('hyperlink'):
                            rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id:
                                try:
                                    hyperlink_url = doc.part.rels[rel_id].target_ref
                                    # Check if this run is within this hyperlink
                                    for run_elem in hyperlink.iter():
                                        if run_elem == run.element:
                                            break
                                    else:
                                        hyperlink_url = None  # This run is not part of this hyperlink
                                        continue
                                    break
                                except (KeyError, AttributeError):
                                    pass  # Invalid relationship, ignore hyperlink
                    
                # Apply Markdown formatting
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                
                # Add space if previous run ended with formatting and this one starts with formatting
                if para_text and para_text[-1] in ['*', ')'] and run_text.startswith(('*', '[')):
                    para_text += " "
                
                para_text += run_text
            
            if para_text.strip():
                # Add list formatting if this is a list item
                if is_list_item:
                    clean_text = para_text.strip()
                    if is_numbered_list:
                        # Preserve numbered list format - remove the original marker and add it back
                        # This ensures consistent spacing while preserving the numbering
                        marker_length = len(list_marker)
                        if clean_text.startswith(list_marker):
                            clean_text = clean_text[marker_length:].strip()
                        para_text = f"{list_marker} {clean_text}"
                    else:
                        # Handle unordered lists - convert to markdown format
                        if clean_text.startswith(('•', '-', '*')):
                            clean_text = clean_text[1:].strip()
                        para_text = f"- {clean_text}"
                
                # Handle spacing for different list types
                if formatted_text and not (para_text.startswith('- ') or is_numbered_list):
                    formatted_text += " "
                elif formatted_text and (para_text.startswith('- ') or is_numbered_list):
                    formatted_text += "\n"
                formatted_text += para_text
        
        return formatted_text.strip()
    
    # Iterate through all tables in the document
    for table in doc.tables:
        # Skip header row (index 0) and process data rows
        for row in table.rows[1:]:  # Skip header row
            cells = row.cells
            if len(cells) >= 3:  # Ensure we have Label, Text, ID columns
                label = cells[0].text.strip()
                # Use formatted text extraction for the text column
                text = _extract_formatted_text_from_cell(cells[1])
                id_value = cells[2].text.strip()
                
                # Map ID to updated text content
                if id_value:  # Only add if ID exists
                    id_to_text[id_value] = text
    
    return id_to_text


def update_csv_with_word_changes(original_csv_data: List[Dict[str, str]], 
                                word_updates: Dict[str, str]) -> List[Dict[str, str]]:
    """Update original CSV data with changes from Word document.
    
    Args:
        original_csv_data: List of dictionaries from original CSV
        word_updates: Dictionary mapping ID to updated text content
        
    Returns:
        Updated CSV data with new text content
    """
    updated_data = []
    
    for row in original_csv_data:
        # Create a copy of the original row
        updated_row = row.copy()
        
        # Get the ID for this row
        row_id = row.get('id', '').strip()
        
        # If we have an update for this ID, apply it
        if row_id in word_updates:
            updated_row['figma_text'] = word_updates[row_id]
        
        updated_data.append(updated_row)
    
    return updated_data


def write_csv_data(csv_data: List[Dict[str, str]], output_path: str) -> None:
    """Write CSV data to file.
    
    Args:
        csv_data: List of dictionaries to write as CSV
        output_path: Path where the CSV file should be saved
    """
    if not csv_data:
        raise ValueError("No data to write to CSV")
    
    # Get all possible fieldnames from the data
    fieldnames = list(csv_data[0].keys())
    
    with open(output_path, 'w', newline='', encoding='utf-8') as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(csv_data)


def extract_word_document_to_csv_format(word_file_path: str, preserve_formatting: bool = True) -> List[Dict[str, str]]:
    """Extract content from a Word document and format it as CSV data.
    
    This function parses a Word document following the expected structure:
    - Headers become both frame and group names
    - Tables following headers contain: ID, Label, Text columns
    - Only outputs: id, frame, group, layer_name, figma_text columns
    
    Args:
        word_file_path: Path to the input Word document
        preserve_formatting: Whether to preserve formatting as Markdown (True) or extract plain text (False)
        
    Returns:
        List of dictionaries with CSV-compatible structure
    """
    doc = Document(word_file_path)
    csv_data = []
    current_frame_group = "General Content"
    
    def _extract_formatted_text_from_cell(cell) -> str:
        """Extract text from a cell while preserving formatting if requested."""
        if not preserve_formatting:
            return cell.text.strip()
        
        formatted_text = ""
        
        for para in cell.paragraphs:
            para_text = ""
            
            # Check if this paragraph is a list item
            is_list_item = False
            is_numbered_list = False
            list_marker = ""
            if para.text.strip():
                text = para.text.strip()
                # Check for unordered list markers
                if text.startswith(('•', '-', '*')):
                    is_list_item = True
                    list_marker = text[0]
                # Check for numbered lists (more comprehensive)
                elif len(text) > 2:
                    # Check for patterns like "1.", "2)", "a.", "A)", "i.", "IV.", etc.
                    numbered_pattern = r'^(\d+[.)]|[a-zA-Z][.)]|[ivxlcdm]+[.)]|[IVXLCDM]+[.)])\s'
                    match = re.match(numbered_pattern, text, re.IGNORECASE)
                    if match:
                        is_list_item = True
                        is_numbered_list = True
                        list_marker = match.group(1)
            
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue
                
                # Check if this run contains a hyperlink
                hyperlink_url = None
                
                # Method 1: Check if run's parent is a hyperlink element
                parent = run.element.getparent()
                while parent is not None:
                    if parent.tag.endswith('hyperlink'):
                        # Found hyperlink element, extract the relationship ID
                        rel_id = parent.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if rel_id:
                            try:
                                hyperlink_url = doc.part.rels[rel_id].target_ref
                                break
                            except (KeyError, AttributeError):
                                pass  # Invalid relationship, ignore hyperlink
                    parent = parent.getparent()
                
                # Method 2: Check for hyperlinks in the paragraph's XML structure
                if not hyperlink_url:
                    for hyperlink in para.element.iter():
                        if hyperlink.tag.endswith('hyperlink'):
                            rel_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                            if rel_id:
                                try:
                                    hyperlink_url = doc.part.rels[rel_id].target_ref
                                    # Check if this run is within this hyperlink
                                    for run_elem in hyperlink.iter():
                                        if run_elem == run.element:
                                            break
                                    else:
                                        hyperlink_url = None  # This run is not part of this hyperlink
                                        continue
                                    break
                                except (KeyError, AttributeError):
                                    pass  # Invalid relationship, ignore hyperlink
                    
                # Apply Markdown formatting
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                
                # Add space if previous run ended with formatting and this one starts with formatting
                if para_text and para_text[-1] in ['*', ')'] and run_text.startswith(('*', '[')):
                    para_text += " "
                
                para_text += run_text
            
            if para_text.strip():
                # Add list formatting if this is a list item
                if is_list_item:
                    clean_text = para_text.strip()
                    if is_numbered_list:
                        # Preserve numbered list format - remove the original marker and add it back
                        # This ensures consistent spacing while preserving the numbering
                        marker_length = len(list_marker)
                        if clean_text.startswith(list_marker):
                            clean_text = clean_text[marker_length:].strip()
                        para_text = f"{list_marker} {clean_text}"
                    else:
                        # Handle unordered lists - convert to markdown format
                        if clean_text.startswith(('•', '-', '*')):
                            clean_text = clean_text[1:].strip()
                        para_text = f"- {clean_text}"
                
                # Handle spacing for different list types
                if formatted_text and not (para_text.startswith('- ') or is_numbered_list):
                    formatted_text += " "
                elif formatted_text and (para_text.startswith('- ') or is_numbered_list):
                    formatted_text += "\n"
                formatted_text += para_text
        
        return formatted_text.strip()
    
    # Process document elements in order
    for element in doc.element.body:
        # Check if element is a paragraph
        if element.tag.endswith('p'):
            # Find the corresponding paragraph object
            for para in doc.paragraphs:
                if para._element == element:
                    # Check if this is a heading
                    style_name = para.style.name.lower()
                    if 'heading' in style_name or 'title' in style_name:
                        text_content = para.text.strip()
                        if text_content:
                            current_frame_group = text_content
                    break
        
        # Check if element is a table
        elif element.tag.endswith('tbl'):
            # Find the corresponding table object
            for table in doc.tables:
                if table._element == element:
                    # Process table rows (skip header row if it exists)
                    for row_idx, row in enumerate(table.rows):
                        cells = row.cells
                        
                        # Expect table structure: Label | Text | ID (3 columns)
                        # Or: ID | Label | Text (3 columns)
                        # Or: Label | Text (2 columns, generate ID)
                        if len(cells) >= 2:
                            # Skip header row - only skip if it's clearly a header row
                            cell_texts = [cell.text.strip().lower() for cell in cells]
                            
                            # Skip if this row contains only header-like words (exact matches for common headers)
                            is_header_row = False
                            if (len(cells) == 3 and 
                                cell_texts[0] in ['label', 'component'] and 
                                cell_texts[1] in ['text', 'description'] and 
                                cell_texts[2] in ['id']):
                                is_header_row = True
                            elif (len(cells) == 2 and 
                                  cell_texts[0] in ['label', 'component'] and 
                                  cell_texts[1] in ['text', 'description']):
                                is_header_row = True
                            
                            if is_header_row:
                                continue
                            
                            # Extract data based on number of columns
                            if len(cells) == 3:
                                # Extract all column texts first
                                col1_text = _extract_formatted_text_from_cell(cells[0])
                                col2_text = _extract_formatted_text_from_cell(cells[1])
                                col3_text = _extract_formatted_text_from_cell(cells[2])
                                
                                # Check if third column looks like an ID (contains special characters like :, ;)
                                if (len(col3_text) > 0 and 
                                    any(char in col3_text for char in [':', ';', 'I2016', 'I-', '_', '-']) and
                                    len(col3_text) > 10):
                                    # Format: Label | Text | ID
                                    layer_name = col1_text
                                    figma_text = col2_text
                                    id_value = col3_text
                                elif (len(col1_text) > 0 and 
                                      any(char in col1_text for char in [':', ';', 'I2016', 'I-', '_', '-']) and
                                      len(col1_text) > 10):
                                    # Format: ID | Label | Text
                                    id_value = col1_text
                                    layer_name = col2_text
                                    figma_text = col3_text
                                else:
                                    # Default: Label | Text | ID (assume third column is ID)
                                    layer_name = col1_text
                                    figma_text = col2_text
                                    id_value = col3_text if col3_text.strip() else f"generated_{len(csv_data) + 1}"
                            
                            elif len(cells) == 2:
                                # Format: Label | Text (generate ID)
                                layer_name = _extract_formatted_text_from_cell(cells[0])
                                figma_text = _extract_formatted_text_from_cell(cells[1])
                                id_value = f"generated_{len(csv_data) + 1}"
                            
                            else:
                                # Single column or more than 3 columns - use first as text
                                layer_name = "Content"
                                figma_text = _extract_formatted_text_from_cell(cells[0])
                                id_value = f"generated_{len(csv_data) + 1}"
                            
                            # Only add row if there's actual content
                            if figma_text.strip():
                                csv_data.append({
                                    "id": id_value,
                                    "frame": current_frame_group,
                                    "group": current_frame_group,
                                    "layer_name": layer_name if layer_name.strip() else "Content",
                                    "figma_text": figma_text
                                })
                    break
    
    return csv_data
